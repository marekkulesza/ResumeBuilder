# pip install python-docx
# pip install docx2pdf
# pip install requests
# pip install beautifulsoup4

# this program needs a resume template doc
# there is a keyword bug that can occur in title_scrape
# if the job title has " - " in it the program will screw up
# Go in and edit your resume and add beside your name Style1 text then save

import requests
from bs4 import BeautifulSoup
import docx
import os
from docx2pdf import convert

#### ONLY EDIT THESE ####

your_name = "Marek Kulesza" # Your Name
resume_path = '/Users/ayylmbo/Desktop/Python/Python Automation Project/Resume.docx' # path to your resume
job_characters = 35  # default 35, low this if you have a long name
new_doc_path = '/Users/ayylmbo/Desktop/Python/Python Automation Project/Marek Kulesza.docx' # path to the new doc so that it doesnt overwrite the resume
pdf_path = '/Users/ayylmbo/Desktop/Resumes/' # path to the completed pdf

#### ONLY EDIT THESE #### 

job_title2 = ""

def title_scrape():
    print("Insert the website address")
    res = requests.get(input())
    # print(res.text) prints the website for debugging
    soup = BeautifulSoup(res.text, 'lxml')
    title_name = soup.select('title')
    title_name2 = (title_name[0].getText())

    keyword = " - " # uses " - " as a seperation for my words and seperates them into 3 lists (before, at " - ", and after)
    after_keyword = title_name2.partition(keyword)
    global job_title2
    job_title = after_keyword[0]
    job_title2 = job_title.title()
    print(job_title2)

def title_changer():
    os.getcwd()
    doc = docx.Document(resume_path) #  this is your Resume template
    print("")
    new_input = job_title2 # What the job title you want 
    if (len(new_input) > job_characters): # the job title is greater than XX characters long, just print out my name
        pass
    else:
        doc.paragraphs[0].add_run(" - "+new_input)

    doc.save(new_doc_path) # Saves to a document which later gets turned into a PDF

def doc_pdf():
    new_doc = docx.Document(new_doc_path)
    print(new_doc.paragraphs[0].text)
    convert(new_doc_path, pdf_path+(new_doc.paragraphs[0].text)+".pdf")

while True:
    title_scrape()
    title_changer()
    doc_pdf()
    print("\nComplete, press enter to add another link or type exit")
    new_input = input()
    if new_input == "exit":
        break